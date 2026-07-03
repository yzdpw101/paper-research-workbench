# -*- coding: utf-8 -*-
"""
python-docx 格式辅助模块 v2 — 增强版：公式 + 交叉引用 + 书签
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter
import os

# ─── MML2OMML XSLT 路径 ───────────────────────────────
MML2OMML_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_omml_xslt = None

def _get_omml_xslt():
    global _omml_xslt
    if _omml_xslt is None and os.path.isfile(MML2OMML_PATH):
        _omml_xslt = etree.parse(MML2OMML_PATH)
    return _omml_xslt

def get_omml_root(latex_expr):
    """LaTeX → MathML → OMML (Office Math Markup Language) XML 元素"""
    # 1) LaTeX → MathML
    mathml_str = latex2mathml.converter.convert(latex_expr)
    mathml_tree = etree.fromstring(mathml_str)
    # 2) MathML → OMML
    xslt = _get_omml_xslt()
    if xslt is not None:
        transform = etree.XSLT(xslt)
        omml_tree = transform(mathml_tree).getroot()
    else:
        omml_tree = mathml_tree  # fallback: use MathML (Word may still render)
    return omml_tree

# ─── 书签与字段码管理器 ────────────────────────────
_bookmark_counter = [0]

def _next_bm_id():
    _bookmark_counter[0] += 1
    return _bookmark_counter[0]

# ─── 字体 ─────────────────────────────────────────────
def set_run_font(run, latin_name="Times New Roman", east_asia_name="宋体", size_pt=12, bold=False, italic=False, color=None):
    run.font.name = latin_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia_name)

# ─── 公式 ─────────────────────────────────────────────
def add_equation(doc, latex_expr, display=True):
    """插入 LaTeX 公式，渲染为 Word 数学公式"""
    omml = get_omml_root(latex_expr)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    if display and omml.tag != qn('m:oMathPara'):
        omath_para = etree.SubElement(etree.Element('dummy'), qn('m:oMathPara'), nsmap={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})
        omath_para.append(omml)
        p._element.append(omml)
    else:
        p._element.append(omml)
    return p

# ─── 书签 ─────────────────────────────────────────────
def add_bookmark_around_element(parent_elem, element, bm_name, bm_id):
    """在 parent_elem 中的 element 前后插入 bookmarkStart/End"""
    idx = list(parent_elem).index(element)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), bm_name)
    parent_elem.insert(idx, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    parent_elem.insert(idx + 2, end)

def add_bookmark_paragraph(paragraph, bm_name):
    """为整个段落添加书签（放在段首）"""
    bm_id = _next_bm_id()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), bm_name)
    paragraph._element.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    paragraph._element.append(end)
    return bm_name

def add_bookmark_around_run(run, bm_name):
    """在 run 周围添加书签"""
    bm_id = _next_bm_id()
    parent = run._element.getparent()
    add_bookmark_around_element(parent, run._element, bm_name, bm_id)
    return bm_name

# ─── 字段码 (field codes) ──────────────────────────
def _insert_complex_field(paragraph, field_code, display_text="?"):
    """
    插入完整字段码：begin + instrText(REF/SEQ) + separate + display + end
    返回 display_run（可加书签）
    """
    rb = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    rb._element.append(fb)
    
    ri = paragraph.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = f" {field_code} "
    ri._element.append(ins)
    
    rs = paragraph.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    rs._element.append(fs)
    
    rd = paragraph.add_run(display_text)
    
    re = paragraph.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    re._element.append(fe)
    
    return rd

def add_ref_field(paragraph, bookmark_name, display_text="?"):
    """插入 REF 字段码（带超链接）"""
    return _insert_complex_field(paragraph, f'REF {bookmark_name} \\h \\n', display_text)

def add_seq_field(paragraph, seq_name, display_text="1"):
    """插入 SEQ 自动编号"""
    return _insert_complex_field(paragraph, f'SEQ {seq_name} \\* ARABIC', display_text)

def add_pageref_field(paragraph, bookmark_name, display_text="?"):
    """插入 PAGEREF 字段码（页码引用）"""
    return _insert_complex_field(paragraph, f'PAGEREF {bookmark_name} \\h', display_text)

# ─── 引用标记 ──────────────────────────────────────
def add_citation_marker(paragraph, ref_number):
    """在正文中插入 [N] 引用标记，可链接到参考文献"""
    rb = paragraph.add_run("[")
    set_run_font(rb, "Times New Roman", "宋体", 12)
    rb.font.superscript = True
    rd = _insert_complex_field(paragraph, f'REF cite_{ref_number} \\h \\* MERGEFORMAT', str(ref_number))
    set_run_font(rd, "Times New Roman", "宋体", 12)
    rd.font.superscript = True
    re = paragraph.add_run("]")
    set_run_font(re, "Times New Roman", "宋体", 12)
    re.font.superscript = True

def add_body_with_citations(doc, segments):
    """
    正文含引用。segments: list of (text, is_citation, ref_num)
    is_citation=True 时插入 citation_marker(ref_num)
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    for seg in segments:
        if len(seg) == 3:
            text, is_citation, ref_num = seg
        else:
            text, is_citation = seg
            ref_num = None
        if is_citation and ref_num:
            add_citation_marker(p, ref_num)
        else:
            run = p.add_run(text)
            set_run_font(run, "Times New Roman", "宋体", 12)

# ─── 标题与正文 ──────────────────────────────────────
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "黑体", 18, bold=True)

def add_abstract_heading(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("摘  要")
    set_run_font(run, "Times New Roman", "黑体", 14, bold=True)

def add_abstract_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "宋体", 12)

def add_keywords(doc, keywords_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    run_label = p.add_run("关键词：")
    set_run_font(run_label, "Times New Roman", "黑体", 12, bold=True)
    run_content = p.add_run(keywords_text)
    set_run_font(run_content, "Times New Roman", "宋体", 12)

def add_heading1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "黑体", 16, bold=True)

def add_heading2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.clear()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "黑体", 14, bold=True)

def add_heading3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.clear()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "黑体", 12, bold=True)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "Times New Roman", "宋体", 12)

def add_body_auto(doc, text):
    """
    正文 + 自动检测 [N] 引用标记并替换为 REF 字段（可点击跳转）
    例如: '文献[1]使用BGA' → '文献' + REF cite_1 + '使用BGA'
    支持: [1] [1,2] [1-3] 格式
    """
    import re
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    
    # 按引用标记分割文本
    parts = re.split(r'(\[\d+(?:[,|-]\d+)*\])', text)
    for part in parts:
        m = re.match(r'\[(\d+(?:[,|-]\d+)*)\]', part)
        if m:
            # 提取第一个数字作为 REF 目标
            nums_str = m.group(1)
            first_num = int(re.match(r'\d+', nums_str).group())
            
            run_b = p.add_run("[")
            set_run_font(run_b, "Times New Roman", "宋体", 12)
            run_b.font.superscript = True
            
            rd = _insert_complex_field(p, f'REF cite_{first_num} \\h \\* MERGEFORMAT', str(first_num))
            set_run_font(rd, "Times New Roman", "宋体", 12)
            rd.font.superscript = True
            
            run_e = p.add_run("]")
            set_run_font(run_e, "Times New Roman", "宋体", 12)
            run_e.font.superscript = True
        else:
            run = p.add_run(part)
            set_run_font(run, "Times New Roman", "宋体", 12)

# ─── 图片与图题（含交叉引用）───────────────────────
_fig_counter = [0]

def add_figure_with_caption(doc, image_path, caption_text, width_cm=14):
    """嵌入图片 + 图题（图号用 SEQ 自动编号 + 书签）"""
    _fig_counter[0] += 1
    bm_name = f"fig_{_fig_counter[0]}"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    
    run_lbl = cap.add_run("图 ")
    set_run_font(run_lbl, "Times New Roman", "宋体", 10.5, bold=True)
    display_run = add_seq_field(cap, "figure", str(_fig_counter[0]))
    set_run_font(display_run, "Times New Roman", "宋体", 10.5)
    run_cap = cap.add_run(f"  {caption_text}")
    set_run_font(run_cap, "Times New Roman", "宋体", 10.5)
    
    add_bookmark_around_run(display_run, bm_name)
    return bm_name

def add_figure_placeholder(doc, caption_text, source_text=None):
    """图片占位：标注来源，图号自动编号"""
    _fig_counter[0] += 1
    bm_name = f"fig_{_fig_counter[0]}"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    placeholder = f"[图片位置：{source_text}]" if source_text else "[图片位置：待插入]"
    run = p.add_run(placeholder)
    set_run_font(run, "Times New Roman", "宋体", 10.5, italic=True, color=RGBColor(128, 128, 128))
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    run_lbl = cap.add_run("图 ")
    set_run_font(run_lbl, "Times New Roman", "宋体", 10.5, bold=True)
    display_run = add_seq_field(cap, "figure", str(_fig_counter[0]))
    add_bookmark_around_run(display_run, bm_name)
    run_cap = cap.add_run(f"  {caption_text}")
    set_run_font(run_cap, "Times New Roman", "宋体", 10.5)
    return bm_name

# ─── 表格与表题 ──────────────────────────────────────
_tab_counter = [0]

def add_table_caption(doc, text):
    _tab_counter[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run_lbl = p.add_run("表 ")
    set_run_font(run_lbl, "Times New Roman", "宋体", 10.5, bold=True)
    display_run = add_seq_field(p, "table", str(_tab_counter[0]))
    run_cap = p.add_run(f"  {text}")
    set_run_font(run_cap, "Times New Roman", "宋体", 10.5)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = 1
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "Times New Roman", "黑体", 10.5, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, "Times New Roman", "宋体", 10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table

# ─── 参考文献（带书签）───────────────────────────────
def add_reference(doc, ref_number, ref_text):
    """参考文献条目，编号 [N] 加书签（可被 REF 引用），书签只包裹数字"""
    bm_name = f"cite_{ref_number}"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    
    # 左括号（不在书签内）
    run_lb = p.add_run("[")
    set_run_font(run_lb, "Times New Roman", "宋体", 10.5)
    
    # 数字（加书签）
    run_num = p.add_run(str(ref_number))
    set_run_font(run_num, "Times New Roman", "宋体", 10.5)
    add_bookmark_around_run(run_num, bm_name)
    
    # 右括号 + 文本（不在书签内）
    run_rb = p.add_run(f"]  {ref_text}")
    set_run_font(run_rb, "Times New Roman", "宋体", 10.5)

# ─── 页面设置与页码 ──────────────────────────────────
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _insert_complex_field(p, "PAGE", "1")

def setup_page(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

# ─── 样式修正与字段自动更新 ─────────────────────────
def configure_heading_styles(doc):
    """
    修正 Heading 1/2/3 样式：黑色文字（去除主题蓝色）+ 中文字体 + 字号
    在 Document() 后立即调用
    """
    configs = {
        'Heading 1': {'size': 16, 'bold': True},
        'Heading 2': {'size': 14, 'bold': True},
        'Heading 3': {'size': 12, 'bold': True},
    }
    for name, cfg in configs.items():
        style = doc.styles[name]
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(cfg['size'])
        style.font.bold = cfg['bold']
        style.font.name = 'Times New Roman'
        # 设置东亚中文字体
        rPr = style.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style.element.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体')

def set_update_fields_on_open(doc):
    """
    设置文档打开时自动更新所有字段（SEQ/REF/PAGE等）
    写入 <w:updateFields w:val="true"/> 到 word/settings.xml
    """
    settings = doc.settings
    update_fields = settings.element.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.element.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

