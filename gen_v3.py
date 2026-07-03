# -*- coding: utf-8 -*-
"""Generate v3 DOCX with fixed heading styles and references"""
import os, sys
sys.path.insert(0, r'E:/Desktop/稀布综述')
sys.path.insert(0, '.')
from docx import Document
from docx_formatter import *

TARGET_DIR = r'E:/Desktop/稀布综述'
FIG_DIR = r'E:/Desktop/稀布综述/figures'
OUTPUT = os.path.join(TARGET_DIR, '阵列天线稀布优化技术综述_v3.docx')

doc = Document()
setup_page(doc)

add_title(doc, '阵列天线稀布优化技术综述')
add_abstract_heading(doc)
add_abstract_body(doc, '阵列天线的稀布优化技术是降低天线系统成本、减轻重量的重要手段。本文系统综述了2020-2026年间阵列天线稀布优化领域的研究进展，覆盖线阵、平面阵、圆环阵、多圆环同心环阵和共形阵五种典型阵型，从群体智能、凸优化、机器学习和解析混合四大类方法进行系统分类总结。')
add_keywords(doc, '阵列天线；稀布优化；稀疏阵列；旁瓣电平抑制；群体智能算法；凸优化；机器学习；共形阵')

add_heading1(doc, '1  引言')
add_body(doc, '阵列天线广泛应用于雷达、通信、电子侦察和射电天文等领域。传统均匀间隔阵列在给定孔径下需要大量阵元，导致系统成本高、重量大。稀布阵列技术通过优化阵元位置和激励分布，在减少有源阵元数量的同时抑制旁瓣电平。近年来，随着5G/6G通信和毫米波雷达的发展，稀布优化技术受到广泛关注。')

add_heading1(doc, '2  稀布阵列天线问题建模')
add_heading2(doc, '2.1  阵列方向图函数')
add_body(doc, '考虑N个各向同性阵元的任意构型阵列，远场方向图函数为：')
add_equation(doc, r'F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta,\phi) \exp[j(k r_n \cos\psi_n + \phi_n)]')
add_body(doc, '等幅同相激励时简化为：')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{n=1}^{N} \exp[jk(x_n\sin\theta\cos\phi + y_n\sin\theta\sin\phi + z_n\cos\theta)]')

add_heading2(doc, '2.2  各类阵型阵列因子')
add_heading3(doc, '2.2.1  线阵')
add_equation(doc, r'\mathrm{AF}(\theta)=\sum_{n=1}^{N} \exp(jk x_n \sin\theta)')
add_heading3(doc, '2.2.2  平面阵')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{m=1}^{M}\sum_{n=1}^{N} I_{mn} \exp[jk(x_m\sin\theta\cos\phi + y_n\sin\theta\sin\phi)]')
add_heading3(doc, '2.2.3  圆环阵')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{q=1}^{Q}\sum_{n=1}^{N_q} \exp[jk R_q (\sin\theta\cos(\phi-\phi_{qn}))]')
add_heading3(doc, '2.2.4  共形阵')
add_equation(doc, r'F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta-\theta_n,\phi-\phi_n) \exp[jk\mathbf{r}_n\cdot\hat{\mathbf{r}}]')

add_heading1(doc, '3  稀布阵列优化方法综述')
add_heading2(doc, '3.1  群体智能算法')
add_heading3(doc, '3.1.1  遗传算法')
add_body(doc, '遗传算法是稀布优化中最经典的方法。文献[1]使用BGA对线阵进行稀疏优化。文献[2]提出基于量子选择机制的GA用于平面阵。文献[3]采用自适应多点变异GA优化大规模阵列。文献[4]研究了线阵稀疏对方向图参数的影响。')
add_heading3(doc, '3.1.2  粒子群优化')
add_body(doc, '文献[5]采用BPSO实现平面阵稀疏综合。文献[6]提出多策略多目标PSO。文献[7]将模拟退火融入PSO设计MIMO稀疏阵列。')
add_heading3(doc, '3.1.3  其他群智能算法')
add_body(doc, '文献[8]提出改进二进制灰狼优化。文献[9]使用二进制蝙蝠群算法。文献[10]将布谷鸟搜索与凸优化结合。文献[11]提出蜣螂优化。文献[12]采用改进蛇优化实现共形阵稀疏。')

f1 = os.path.join(FIG_DIR, '算法/9753378-fig-3-source.gif')
if os.path.isfile(f1):
    add_figure_with_caption(doc, f1, '基于CNN的阵列天线稀疏综合流程（文献[13]）')

add_heading2(doc, '3.2  凸优化与压缩感知')
add_body(doc, '文献[14]通过迭代凸优化综合稀疏阵列。文献[15]将ADMM应用于方向图综合。文献[16]采用共识ADMM实现多波束综合。文献[17]提出重加权l1范数最小化。文献[18]基于离网压缩感知实现低复杂度综合。文献[19]结合贝叶斯CS与阵列扩张。文献[20]将分支定界与凸优化结合。')

add_heading2(doc, '3.3  机器学习方法')
add_body(doc, '文献[13]使用CNN实现阵列稀疏综合。文献[21]提出DNN线阵稀疏方法。文献[22]通过强化学习优化概率密度锥削分布。')

add_heading2(doc, '3.4  解析与混合方法')
add_body(doc, '文献[23]使用两级IFT实现稀疏平面阵。文献[24]将JAYA与IFT混合。文献[25]提出OMP-Broyden混合算法。文献[26]采用量子-经典混合进化优化。')

add_table_caption(doc, '各类稀布优化方法对比')
add_table(doc, ['方法类别','代表算法','优势','局限性','适用阵型'],
    [['群体智能','GA/BPSO/GWO/CS','无梯度全局搜索','计算量大','全部阵型'],
     ['凸优化/CS','ADMM/SOCP','收敛保证','约束敏感','线阵/平面阵'],
     ['机器学习','CNN/DNN/RL','推理快','泛化待验证','线阵/平面阵'],
     ['解析方法','IFT/矩阵束','计算极快','灵活性差','线阵/平面阵'],
     ['混合方法','JAYA-IFT/Broyden','兼顾效率与搜索','设计复杂','线阵/平面阵']])

add_heading1(doc, '4  各阵型稀布优化研究进展')
add_heading2(doc, '4.1  稀布线阵')
add_body(doc, '线阵的稀布优化研究最为成熟。文献[27]提出基于PSF的适应度函数。文献[28]针对太赫兹频段进行GA稀疏设计。文献[29]提出高效综合技术。文献[30]引入低比特相位量化约束。文献[31]综合比较多种群智能算法性能。')

f2 = os.path.join(FIG_DIR, '线阵/9964507-fig-4-source.gif')
if os.path.isfile(f2):
    add_figure_with_caption(doc, f2, 'BGA优化后稀布线阵的方向图（文献[1]）')

add_heading2(doc, '4.2  稀布平面阵')
add_body(doc, '文献[32]提出预编码子阵结构。文献[33]优化等幅激励平面阵增益。文献[34]采用混合无约束-启发式方法。')

f3 = os.path.join(FIG_DIR, '平面阵/yang4-3088492.gif')
if os.path.isfile(f3):
    add_figure_with_caption(doc, f3, '迭代凸优化稀疏平面阵3D方向图（文献[14]）')

add_heading2(doc, '4.3  圆环阵与多圆环同心环阵')
add_body(doc, '文献[35]设计频率不变波束形成稀疏同心圆环阵。文献[36]提出加窗波束形成器。文献[37]研究低旁瓣稀疏圆环阵。文献[38]利用贪婪算法设计宽带MVDR同心圆环阵。')

f4 = os.path.join(FIG_DIR, '圆环阵/wang1-p5-wang.gif')
if os.path.isfile(f4):
    add_figure_with_caption(doc, f4, '稀疏圆环阵阵元分布与方向图（文献[37]）')

add_heading2(doc, '4.4  共形阵')
add_body(doc, '文献[39]实现稀疏球面共形阵。文献[40]设计导航共形阵。文献[41]利用AEP方法实现低旁瓣。文献[42]基于多智能体GA。文献[43]设计多面阵子阵结构。文献[44][45]分别基于Sobol和Van der Corput序列实现柱面阵低偏差采样。文献[46]通过多目标混合方法优化方向图。')

f5 = os.path.join(FIG_DIR, '共形阵/10069223-fig-2-source.gif')
if os.path.isfile(f5):
    add_figure_with_caption(doc, f5, '多智能体GA优化的稀疏共形阵阵元分布（文献[42]）')

add_heading1(doc, '5  工程约束与讨论')
add_heading2(doc, '5.1  互耦效应')
add_body(doc, '稀布阵列的非均匀间距使互耦效应复杂化。文献[41]通过引入有源单元方向图（AEP）考虑互耦影响。')
add_heading2(doc, '5.2  幅相误差与量化约束')
add_body(doc, '文献[30]将低比特相位量化约束嵌入优化模型。文献[17]采用混合整数规划处理离散相位约束。')
add_heading2(doc, '5.3  宽带特性')
add_body(doc, '文献[35]针对宽带波束形成设计频率不变稀疏圆环阵。文献[38]设计宽带MVDR波束形成的同心圆环阵。')

add_heading1(doc, '6  未来展望')
add_body(doc, '（1）AI驱动的端到端设计：深度学习与电磁仿真深度耦合实现全流程自动化。')
add_body(doc, '（2）亚波长与可重构稀布阵列：结合超表面和MEMS开关实现波束可重构。')
add_body(doc, '（3）MIMO雷达与通信一体化：大规模MIMO需同时考虑通信和感知性能。')
add_body(doc, '（4）多物理场协同优化：结构、热管理和电磁性能的联合优化。')
add_body(doc, '（5）量子计算应用：量子-经典混合算法为大规模阵列优化提供新路径。')

add_heading1(doc, '参考文献')
refs = [
    'M. M. Kamal et al., "Optimization of Linear Antenna Array Thinning using BGA," ICCCI 2022.',
    '"Planar Array Thinning by GA with Quantum Selection," 2025.',
    '"Sparse Antenna Array Synthesis via Adaptive Multipoint Mutation GA," 2025.',
    'A. S. Karasev et al., "Effect of Linear Array Thinning on Pattern Parameters," EDM 2022.',
    'S. Vankayalapati et al., "BPSO for Thinned Planar Array Synthesis," ICACCS 2022.',
    '"Multi-Objective PSO for Sparse Planar Array Synthesis," 2025.',
    '"Optimal MIMO Sparse Array via SA-PSO," 2022.',
    'L. Su et al., "Linear Array Thinning Based on Improved Binary GWO," ICMMT 2025.',
    'S. Vankayalapati et al., "Binary bat swarm for thinning linear array," ICCCI 2023.',
    'R.-Q. Wang et al., "Hybrid Cuckoo Search with Convex Programming for Sparse Linear Arrays," IEEE AWPL, 2020.',
    '"Planar Sparse Arrays Based on Improved Dung Beetle Optimization," 2025.',
    '"Improved Snake Optimization for Sparse Conformal Array Beamforming," 2024.',
    'S. Shen et al., "Thinned Array Antenna with CNN," ISAPE 2021.',
    '"Sparse Array Synthesis via Iterative Convex Optimization," 2021.',
    '"Sparse Array Beampattern Synthesis via ADMM," 2021.',
    '"Multibeam Synthesis via Consensus PDD Framework," 2023.',
    '"Sparse Arrays With Reweighted l1-norm Minimization," 2022.',
    '"Off-Grid Compressive Sensing for Sparse Array Synthesis," 2022.',
    '"BCS and Array Dilation for Sparse Linear Arrays," 2022.',
    '"Large-Scale Sparse Array via Branch and Bound with Convex Optimization," 2025.',
    '"Linear Sparse Array Using DNN," 2023.',
    '"Adaptive Probability Density Taper for Large Planar Array Thinning," 2020.',
    '"Thinned Planar Array Using Two-Stage IFT," 2025.',
    'X. Yang et al., "Linear Array Thinning Using JAYA-IFT," ACES-China 2023.',
    '"Hybrid OMP-Broyden for Thinning Antenna Array," 2025.',
    '"Array Thinning by Hybrid Quantum-Classical Optimization," 2025.',
    '"New Fitness Function for Sparse Linear Array Based on PSF," 2021.',
    '"Static Thinning of Large Linear Arrays Using GA for THz," 2025.',
    'Z. Zhou et al., "Effective Synthesis for Linear Array Thinning," Radar 2021.',
    '"Sparse Array Optimization with Low-Bit Phase Quantization," 2025.',
    '"Pattern Synthesis Using Mayfly Algorithm," IEEE Access, 2021.',
    '"Thinned Planar Arrays Based on Precoded Subarray Structures," 2022.',
    '"High Gain Optimization for Sparse Planar Array," 2022.',
    '"Hybrid Unconstrained-Heuristic for Sparse Planar Arrays," 2022.',
    'Y. Buchris et al., "Frequency-Invariant Beamformers with Sparse Concentric Circular Arrays," WASPAA 2023.',
    '"Window Beamformer for Sparse Concentric Circular Array," 2021.',
    'Y. Wang et al., "Sparse Circular Arrays with Low Sidelobe Levels," CTISC 2022.',
    '"Greedy Design of Circular Concentric Arrays for Broadband MVDR," 2024.',
    'G. Yang et al., "Optimized Design of Sparse Spherical Conformal Array," APCAP 2024.',
    '"Sparse Design of Navigation Conformal Array," 2025.',
    '"Low Sidelobe Optimization of Conformal Sparse Array using Active Pattern," 2024.',
    'G. Liu et al., "Sparse Conformal Array Synthesis Based on Multiagent GA," APCAP 2022.',
    '"Sparse Conformal Multi-Faceted Array," 2023.',
    '"Sparse Cylindrical Arrays Based on Sobol Sequence Sampling," 2022.',
    '"Sparse Cylindrical Arrays with Van der Corput Sequence," 2021.',
    'C. Liu et al., "Multi-Objective Hybrid for Sparse Conformal Array," ICCCA 2025.',
]
for i, ref in enumerate(refs, 1):
    add_reference(doc, i, ref)

add_page_numbers(doc)
doc.save(OUTPUT)
print(f'Done! Saved to: {OUTPUT}')
print(f'Size: {os.path.getsize(OUTPUT)} bytes')
