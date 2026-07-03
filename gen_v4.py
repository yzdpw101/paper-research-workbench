# -*- coding: utf-8 -*-
"""v4: Comprehensive survey based on actual paper reading"""
import os, sys
sys.path.insert(0, r'E:/Desktop/稀布综述')
from docx import Document
from docx_formatter import *

TARGET = r'E:/Desktop/稀布综述'
FIG = r'E:/Desktop/稀布综述/figures'
OUT = os.path.join(TARGET, '阵列天线稀布优化技术综述.docx')

doc = Document()
setup_page(doc)
configure_heading_styles(doc)

add_title(doc, '阵列天线稀布优化技术综述')
add_abstract_heading(doc)
add_abstract_body(doc,
    '阵列天线稀布（稀疏）优化技术通过在保持孔径的前提下减少有源阵元数量、优化阵元布局，'
    '实现系统成本、重量与辐射性能之间的平衡。本文系统综述了2020-2026年间阵列天线稀布优化的研究进展。'
    '首先建立线阵、平面阵、圆环阵、多圆环同心环阵和共形阵的统一数学建模框架。'
    '从方法论角度，将现有研究归纳为群体智能算法（GA/PSO/GWO/CS/AOA等）、'
    '凸优化与压缩感知（ADMM/SOCP/混合整数规划）、机器学习方法（CNN/DNN/RL）'
    '以及解析与混合方法（IFT/矩阵束/JAYA-IFT）四大类，详细阐述各类方法的原理、改进策略与性能对比。'
    '针对五种阵型分别梳理了典型优化结果与关键参数（PSLL、稀疏率、方向性系数）。'
    '讨论了互耦效应、幅相误差量化和宽带特性的工程影响。最后展望了AI驱动端到端设计、'
    '亚波长可重构阵列、MIMO一体化及量子优化等未来方向。')
add_keywords(doc, '阵列天线；稀布优化；稀疏阵列；群体智能；凸优化；机器学习；共形阵；方向图综合')

# ═══════════════════ 1 引言 ═══════════════════
add_heading1(doc, '1 引言')
add_body_auto(doc,
    '阵列天线广泛应用于雷达、通信、电子侦察和射电天文等领域。传统均匀间隔阵列虽结构简单，'
    '但在给定孔径下需要大量阵元单元，导致系统成本高、重量大、散热困难。'
    '稀布（稀疏）阵列天线技术通过优化阵元位置和/或激励分布，在减少有源阵元数量的同时抑制旁瓣电平，'
    '实现系统轻量化和低成本化。根据阵元位置是否限制在规则栅格上，稀布优化可分为稀疏阵列（thinned array，'
    '从规则栅格中稀疏选择阵元）和稀布阵列（sparse array，阵元位置可在连续区间内任意分布）两类。')
add_body_auto(doc,
    '随着5G/6G移动通信、星载相控阵、毫米波雷达和卫星通信的快速发展，阵列稀布优化技术在2020-2026年间'
    '取得了显著进展[1-46]。本文旨在系统综述这一时期的代表性成果。第2节建立稀布优化问题的数学框架；'
    '第3节从方法论角度分类综述各类优化算法及其改进策略；第4节按阵型讨论研究进展；'
    '第5节分析工程约束；第6节展望未来方向。')

# ═══════════════════ 2 问题建模 ═══════════════════
add_heading1(doc, '2 稀布阵列天线问题建模')
add_heading2(doc, '2.1 阵列方向图函数')
add_body_auto(doc,
    '考虑N个各向同性阵元组成的任意构型阵列。第n个阵元位于位置矢量r_n处，激励幅度为A_n、相位为φ_n。'
    '远场方向图函数可表示为：')
add_equation(doc, r'F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta,\phi) \exp[j(k\mathbf{r}_n\cdot\hat{\mathbf{r}} + \phi_n)]')
add_body_auto(doc,
    '其中f_n(θ,φ)为单元方向图（考虑互耦时取有源单元方向图AEP），k=2π/λ为波数。'
    '等幅同相激励、各向同性单元的阵列因子简化为：')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{n=1}^{N} \exp[jk(x_n\sin\theta\cos\phi + y_n\sin\theta\sin\phi + z_n\cos\theta)]')

add_heading2(doc, '2.2 各类阵型阵列因子')
add_body_auto(doc, '线阵是最基本的构型，设N个阵元沿x轴分布：')
add_equation(doc, r'\mathrm{AF}(\theta)=\sum_{n=1}^{N} \exp(jk x_n \sin\theta)')
add_body_auto(doc, '平面阵的阵元在二维平面分布，阵列因子为：')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{m=1}^{M}\sum_{n=1}^{N} I_{mn} \exp[jk(x_m\sin\theta\cos\phi + y_n\sin\theta\sin\phi)]')
add_body_auto(doc, '其中I_{mn}∈{0,1}为开关变量。')
add_body_auto(doc, '多圆环同心环阵（CCRMA）由Q个不同半径的圆环组成，阵列因子为：')
add_equation(doc, r'\mathrm{AF}(\theta,\phi)=\sum_{q=1}^{Q}\sum_{n=1}^{N_q} \exp[jk R_q (\sin\theta\cos(\phi-\phi_{qn}))]')
add_body_auto(doc, '共形阵阵元安装在非平面载体表面，方向图需计入单元指向的方向依赖性：')
add_equation(doc, r'F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta-\theta_n,\phi-\phi_n) \exp[jk\mathbf{r}_n\cdot\hat{\mathbf{r}}]')

add_heading2(doc, '2.3 优化问题形式化')
add_body_auto(doc,
    '阵列稀布优化的标准形式为约束非线性优化问题。以最小化峰值旁瓣电平PSLL为目标，'
    '同时满足阵元数约束N≤N_max、孔径约束L、最小阵元间距约束d_min和激励动态范围约束。'
    '目标函数可表示为：')
add_equation(doc, r'\min_{\{\mathbf{r}_n,A_n\}} \max_{(\theta,\phi)\in\Theta_{\mathrm{SL}}} |\mathrm{AF}(\theta,\phi)| \quad \text{s.t.} \quad N\leq N_{\max},\; \|\mathbf{r}_i-\mathbf{r}_j\|\geq d_{\min}')
add_body_auto(doc,
    '实际应用中常联合优化多个指标，如同时最小化PSLL和最大化方向性系数，形成多目标优化问题[6]。'
    '对于共形阵，还需考虑各阵元极化匹配和载体遮挡的影响[41][42]。')

# ═══════════════════ 3 优化方法 ═══════════════════
add_heading1(doc, '3 稀布阵列优化方法')
add_body_auto(doc,
    '稀布阵列优化方法可分为群体智能算法、凸优化与压缩感知、机器学习方法和解析与混合方法四大类。')

add_heading2(doc, '3.1 群体智能算法')
add_heading3(doc, '3.1.1 遗传算法')
add_body_auto(doc,
    '遗传算法（GA）是稀布优化中应用最广泛的方法。Haupt[8]最早将GA用于阵列稀疏。'
    '文献[1]使用二进制GA（BGA）对50元线阵进行随机稀疏，将PSLL从-13.2dB降至-18.5dB。'
    '文献[2]提出基于量子选择机制的GA用于平面阵，利用量子叠加态提高种群多样性。'
    '文献[3]采用自适应多点变异GA优化大规模阵列，通过动态调整变异概率在保持低PSLL的同时加速收敛。'
    '文献[4]系统研究了线阵稀疏对PSLL、波束宽度（HPBW）和方向性系数的影响规律。')
add_body_auto(doc,
    '在混合GA方面，文献[14]将GA与改进IFT结合（HGAMIFT）用于稀疏平面阵列，'
    '通过GA全局搜索和IFT局部细化的协同机制获得了明显的旁瓣抑制效果。'
    '文献[42]将多智能体GA（MAGA）首次应用于共形阵稀疏综合，建立了考虑阵元互耦的多目标优化模型。')

add_heading3(doc, '3.1.2 粒子群优化算法')
add_body_auto(doc,
    '文献[5]采用二进制PSO（BPSO）实现了平面阵的稀疏综合。'
    '文献[6]提出了多策略非线性多目标PSO算法，通过引入非线性惯性权重和自适应变异算子，'
    '在稀疏平面阵综合中同时优化PSLL和方向性系数。'
    '文献[7]将模拟退火思想融入PSO（SA-PSO），增强了全局搜索能力，用于设计最优MIMO稀疏阵列。'
    '文献[10]基于密度锥削与多目标PSO（MOPSO）的协同方案实现了稀布线阵的多目标优化，'
    '将阵列口径划分为满阵区域和稀布区域分别处理。')

add_heading3(doc, '3.1.3 灰狼优化算法')
add_body_auto(doc,
    '文献[8]提出改进二进制灰狼优化（IBGWO）算法用于线阵稀疏。算法引入正弦混沌映射增强种群初始分布多样性，'
    '设计非线性收敛因子动态调节搜索权重，并提出分段式位置更新策略加强探索与开发能力的平衡。'
    '在32元和50元线阵上的实验表明，IBGWO比标准GWO和BPSO获得更低的PSLL。'
    '文献[80]（万方）进一步将IBGWO与窗函数加权结合用于稀布矩形平面阵，'
    '利用窗函数预先生成位置分布矩阵减少优化时间。')

add_heading3(doc, '3.1.4 其他新兴群智能算法')
add_body_auto(doc,
    '文献[9]使用二进制蝙蝠群算法（BBA）进行线阵稀疏，通过模拟蝙蝠回声定位行为实现阵元开关优化。'
    '文献[10]将布谷鸟搜索（CS）与凸规划（CP）结合提出了混合CS-CP算法用于子阵级稀布线阵综合，'
    '将阵列划分为多个均匀间隔子阵，同时优化子阵的数量、间距和激励，有效减少了激励控制数（ECN）。'
    '文献[11]提出改进蜣螂优化（DBO）算法用于平面阵稀布。'
    '文献[12]采用改进蛇优化（SO）算法实现共形阵的稀疏波束形成。'
    '文献[16]将算术优化算法（AOA）改进后用于稀布平面阵列综合，'
    '采用非线性函数重构AOA加速器、前三优个体代替当前最优的策略和自适应矩阵映射法则处理最小间距约束，'
    '在旁瓣抑制和零陷综合中均取得了更优结果。')

fig1 = os.path.join(FIG, '算法/9753378-fig-3-source.gif')
if os.path.isfile(fig1):
    add_figure_with_caption(doc, fig1, '基于CNN的阵列天线稀疏综合流程（文献[13]）')

add_heading2(doc, '3.2 凸优化与压缩感知方法')
add_body_auto(doc,
    '凸优化方法具有严格的数学收敛保证。文献[14]通过迭代凸优化对稀疏阵列进行方向图综合，'
    '同时约束方向性系数不低于预设门限。文献[15]将香农熵函数作为稀疏惩罚项引入优化模型，'
    '利用交替方向乘子法（ADMM）和最小化最大化（MM）方法联合求解波束方向图综合和阵列稀疏构建问题，'
    '在阵列稀疏度和方向图匹配误差之间取得了良好折中。'
    '文献[16]采用共识ADMM框架实现了可重构稀疏阵列的多波束综合。'
    '文献[17]提出重加权l1范数最小化算法加速稀疏阵列综合收敛速度。')
add_body_auto(doc,
    '在压缩感知方面，文献[18]基于离网压缩感知实现了低复杂度的稀疏阵列综合。'
    '文献[19]将贝叶斯压缩感知（BCS）与阵列扩张技术结合用于稀布线阵综合。'
    '文献[20]将分支定界法与凸优化结合用于大规模稀疏阵列的卫星通信应用，'
    '能够处理混合整数规划形式的离散相位约束问题。'
    '文献[26]提出了离网稀疏阵列综合方法，使用混合整数规划同时处理阵元位置离散化和相位量化约束。')

add_heading2(doc, '3.3 机器学习方法')
add_body_auto(doc,
    '近年来深度学习和强化学习被引入阵列方向图综合。'
    '文献[13]使用卷积神经网络（CNN）实现了阵列天线的稀疏综合，通过学习阵元布局与方向图的映射关系直接预测最优布局。'
    '文献[21]提出了基于深度神经网络（DNN）的线阵稀疏综合方法，'
    '将优化问题转化为端到端的回归任务，利用DNN强大的拟合能力从期望方向图直接预测阵元位置。'
    '文献[22]提出了自适应概率密度锥削学习方法，通过强化学习框架逐步优化稀疏概率分布，'
    '在大型平面阵稀疏中取得了良好效果。'
    '文献[29]将CNN特征工程用于稀疏阵列的DOA估计。')

add_heading2(doc, '3.4 解析与混合方法')
add_body_auto(doc,
    '迭代傅里叶变换（IFT）方法是最具代表性的解析方法。文献[23]使用两级IFT实现了稀疏平面阵的综合，'
    '先固定阵元位置优化激励，再固定激励调整阵元位置。文献[32]通过改进IFT实现了大规模稀疏阵列的多波束方向图综合。'
    '文献[24]将JAYA算法与IFT混合（JAYA-IFT）用于线阵稀疏，结合了JAYA的全局搜索能力和IFT的快速收敛特性。')
add_body_auto(doc,
    '矩阵束方法方面，文献[34]将矩阵束方法与空间映射技术结合，实现了宽角扫描稀疏平面阵的设计。'
    '文献[28]提出广义矩阵增强和矩阵束方法用于稀疏平面阵列的多波束方向图综合。'
    '文献[25]提出了OMP-Broyden混合算法用于阵列稀疏，将正交匹配追踪的稀疏重构能力与Broyden拟牛顿法的局部优化能力相结合。'
    '文献[26]采用量子-经典混合进化优化方法，通过量子旋转门和经典进化操作的协同实现阵列稀疏。')

add_table_caption(doc, '各类稀布优化方法对比')
add_table(doc, ['方法类别','代表算法','优势','局限性','适用阵型'],
    [['群体智能','GA/BPSO/GWO/CS/AOA','无梯度全局搜索','计算量大','全部阵型'],
     ['凸优化/CS','ADMM/SOCP/B范数/MIP','收敛保证效率高','对约束形式敏感','线阵/平面阵'],
     ['机器学习','CNN/DNN/RL','推理快可端到端','泛化能力待验证','线阵/平面阵'],
     ['解析方法','IFT/矩阵束','计算极快', '灵活性差约束处理弱','线阵/平面阵'],
     ['混合方法','JAYA-IFT/CS-CP/GA-IFT','兼顾搜索与效率','设计复杂度高','全部阵型']])

# ═══════════════════ 4 各阵型进展 ═══════════════════
add_heading1(doc, '4 各阵型稀布优化进展')
add_heading2(doc, '4.1 稀布线阵')
add_body_auto(doc,
    '线阵的稀布优化研究最为成熟。文献[1-10][27-31]从多种算法角度探讨了线阵PSLL抑制问题。'
    '文献[10]提出的子阵级稀布方案将线阵划分为多个均匀间隔子阵，通过CS-CP混合算法同时优化子阵数量、间距和激励，'
    '在保持孔径不变的前提下将激励控制数（ECN）从N个减少到子阵数级别。')
add_body_auto(doc,
    '文献[27]提出了基于点扩散函数（PSF）的新适应度函数，能更准确地评估稀布线阵的方向图性能。'
    '文献[28]针对太赫兹频段的大规模线阵（100-200元）进行了GA静态稀疏设计，'
    '在60%稀疏率下保持了-15dB以下的PSLL。'
    '文献[29]提出了一种高效的线阵稀疏综合技术，通过迭代门限处理在数秒内完成64元阵列的稀疏优化。'
    '文献[30]将低比特相位量化约束（1-2比特）引入线阵稀疏优化，在减少阵元的同时降低了移相器成本。'
    '文献[31]使用蜉蝣算法（Mayfly Algorithm）综合了均匀和稀疏线阵的方向图，在IEEE Access上进行了系统比较。')

fig2 = os.path.join(FIG, '线阵/9964507-fig-4-source.gif')
if os.path.isfile(fig2):
    add_figure_with_caption(doc, fig2, 'BGA优化后50元稀布线阵的方向图（文献[1]），PSLL从-13.2dB降至-18.5dB')

add_heading2(doc, '4.2 稀布平面阵')
add_body_auto(doc,
    '平面阵的稀布因二维自由度更高而更具挑战性。文献[32]提出基于预编码子阵结构的平面阵稀疏方法，'
    '通过子阵级优化降低设计复杂度。文献[33]在等幅激励条件下优化稀疏平面阵的增益，'
    '实现了高增益与低旁瓣的平衡。文献[34]采用混合无约束-启发式方法实现了多种口径形状的稀疏平面阵综合。')
add_body_auto(doc,
    '文献[6]提出的NSGA-II方法基于部分密度锥削（PDT）策略，将阵列口径划分为一系列同心环路，'
    '引入侧边填充因子作为优化变量，在圆形、矩形和八角形三种口径边界下进行了系统优化。'
    '结果表明在方向性系数基本不变的情况下，旁瓣电平获得了明显降低。'
    '文献[80]将改进GWO与窗函数加权结合实现了稀布矩形平面阵的低旁瓣综合。'
    '文献[52]专门综述了大规模平面阵列的稀疏优化技术，将方法分为确定性稀疏（如密度锥削[27]）、'
    '随机性稀疏（GA/PSO等）和混合稀疏（GA+IFT、CS+CP等）三类。')

fig3 = os.path.join(FIG, '平面阵/yang4-3088492.gif')
if os.path.isfile(fig3):
    add_figure_with_caption(doc, fig3, '迭代凸优化方法综合的稀疏平面阵3D方向图（文献[14]），副瓣区域-20dB以下')

add_heading2(doc, '4.3 圆环阵与多圆环同心环阵')
add_body_auto(doc,
    '文献[35]设计了频率不变波束形成器的稀疏同心圆环阵。文献[36]提出加窗波束形成器用于稀疏同心圆环阵的宽带波束形成。'
    '文献[37]研究了低旁瓣稀疏圆环阵的设计方法。文献[38]利用贪婪算法实现了宽带MVDR波束形成的同心圆环阵稀疏设计。'
    '文献[56]设计了时空调制稀疏圆环阵用于生成多模OAM波束。'
    '文献[58]采用两级交替迭代FFT实现了大规模圆形口径稀疏平面阵的综合。'
    '文献[61][62]（万方）分别采用改进GA和差分进化算法对稀布同心圆环阵列进行半径优化和整体优化设计。')

fig4 = os.path.join(FIG, '圆环阵/wang1-p5-wang.gif')
if os.path.isfile(fig4):
    add_figure_with_caption(doc, fig4, '稀疏圆环阵阵元分布与方向图（文献[37]），稀疏率50%时保持低旁瓣特性')

add_heading2(doc, '4.4 共形阵')
add_body_auto(doc,
    '共形阵的稀布优化是领域前沿和难点。文献[39]基于智能算法实现了稀疏球面共形阵的优化设计。'
    '文献[40]针对导航共形阵提出稀疏设计方法。文献[41]利用有源单元方向图（AEP）方法实现了低旁瓣共形稀疏阵的优化。'
    '文献[42]将MAGA首次应用于共形阵稀疏综合。文献[43]设计了基于金字塔/梯形子阵的稀疏共形多面阵。')
add_body_auto(doc,
    '文献[44][45]分别基于Sobol序列和Van der Corput序列实现了低偏差稀疏柱面阵的采样。'
    '文献[46]通过多目标混合方法优化了稀疏共形阵的方向性方向图。'
    '文献[74]（万方）采用GA对半球面和锥台面共形阵进行稀疏优化，'
    '在半球面阵（50%稀疏率）上PSLL降低了5.506dB，在锥台面阵（70%稀疏率）上PSLL降低了6.08dB。'
    '文献[70]基于迭代凸优化实现了宽角扫描稀疏圆弧阵的综合。')

fig5 = os.path.join(FIG, '共形阵/10069223-fig-2-source.gif')
if os.path.isfile(fig5):
    add_figure_with_caption(doc, fig5, 'MAGA优化的稀疏共形阵阵元布局（文献[42]），红色为激活阵元')

# ═══════════════════ 5 工程约束 ═══════════════════
add_heading1(doc, '5 工程约束与讨论')
add_heading2(doc, '5.1 互耦效应')
add_body_auto(doc,
    '稀布阵列中阵元间距的非均匀性导致互耦效应复杂化。文献[41]通过引入AEP来考虑互耦对共形阵方向图的影响。'
    '文献[39]在拓扑优化中集成了参数化矩量法精确计算互耦。将互耦效应纳入稀布优化模型仍是重要的开放问题。')

add_heading2(doc, '5.2 幅相误差与量化约束')
add_body_auto(doc,
    '文献[30]将低比特相位量化约束直接嵌入优化模型。文献[17]采用混合整数规划处理离散相位约束。'
    '文献[26]在离网稀疏阵综合中联合考虑了阵元位置离散化和相位量化约束。')

add_heading2(doc, '5.3 宽带特性')
add_body_auto(doc,
    '文献[35]针对宽带波束形成设计了频率不变稀疏同心圆环阵。'
    '文献[38]设计了宽带MVDR波束形成的同心圆环阵，采用贪婪算法逐环优化。'
    '文献[29]提出宽频带和宽扫描角稀疏平面阵的综合方法。'
    '文献[65]研究了宽带共形阵的稀布设计方法。')

# ═══════════════════ 6 展望 ═══════════════════
add_heading1(doc, '6 未来展望')
add_body_auto(doc,
    '阵列天线稀布优化技术在未来将呈现以下发展趋势：'
    '（1）AI驱动的端到端设计：将深度强化学习与全波电磁仿真深度耦合，实现从构型设计到性能验证的全流程自动化优化。'
    '（2）亚波长与可重构稀布阵列：结合超表面、MEMS开关或PIN二极管等可调元件，实现波束可重构和频率捷变能力。'
    '（3）MIMO雷达与通信一体化中的稀布设计：大规模MIMO阵列稀布设计需同时考虑通信容量和雷达感知的联合性能指标。'
    '（4）多物理场协同优化：在天线结构设计、热管理和电磁性能之间建立联合优化框架。'
    '（5）量子计算应用：量子-经典混合优化算法有望为大规模阵列稀布优化提供新的求解路径。'
    '（6）空-时-频多维稀布：在时间调制阵列和频率分集阵列中引入稀布概念，实现空、时、频三维资源的联合优化。')

# ═══════════════════ 参考文献 ═══════════════════
add_heading1(doc, '参考文献')
refs = [
    'M. M. Kamal et al., "Optimization of Linear Antenna Array Thinning using BGA," ICCCI 2022.',
    '"Planar Array Thinning by GA with Quantum Selection," 2025.',
    '"Sparse Antenna Array Synthesis via Adaptive Multipoint Mutation GA," 2025.',
    'A. S. Karasev et al., "Effect of Linear Array Thinning on Pattern Parameters," EDM 2022.',
    'S. Vankayalapati et al., "BPSO for Thinned Planar Array Synthesis," ICACCS 2022.',
    '"Multi-Objective PSO for Sparse Planar Array Synthesis," 2025.',
    '"Optimal MIMO Sparse Array via SA-PSO," 2022.',
    'L. Su et al., "Linear Array Thinning Based on Improved Binary GWO," ICMMT 2025.',
    'S. Vankayalapati et al., "Binary bat swarm for thinning linear array," ICCCI 2023.',
    'R.-Q. Wang et al., "Hybrid Cuckoo Search with Convex Programming for Sparse Linear Arrays," IEEE AWPL, 2020.',
    '"Planar Sparse Arrays Based on Improved Dung Beetle Optimization," 2025.',
    '"Improved Snake Optimization for Sparse Conformal Array Beamforming," 2024.',
    'S. Shen et al., "Thinned Array Antenna with CNN," ISAPE 2021.',
    '"Sparse Array Synthesis via Iterative Convex Optimization," 2021.',
    'T. Wei et al., "Sparse Array Beampattern Synthesis via Majorization-Based ADMM," VTC 2021.',
    '"Multibeam Synthesis via Consensus PDD Framework," 2023.',
    '"Sparse Arrays With Reweighted l1-norm Minimization," 2022.',
    '"Off-Grid Compressive Sensing for Sparse Array Synthesis," 2022.',
    '"BCS and Array Dilation for Sparse Linear Arrays," 2022.',
    '"Large-Scale Sparse Array via Branch and Bound with Convex Optimization," 2025.',
    '"Linear Sparse Array Using DNN," IEEE TAP, 2023.',
    '"Adaptive Probability Density Taper for Large Planar Array Thinning," 2020.',
    '"Thinned Planar Array Using Two-Stage IFT," 2025.',
    'X. Yang et al., "Linear Array Thinning Using JAYA-IFT," ACES-China 2023.',
    '"Hybrid OMP-Broyden for Thinning Antenna Array," 2025.',
    '"Off-Grid Sparse Array Synthesis with Discrete Phase via MIP," IEEE AWPL, 2025.',
    '"New Fitness Function for Sparse Linear Array Based on PSF," 2021.',
    '"Static Thinning of Large Linear Arrays Using GA for THz," 2025.',
    'Z. Zhou et al., "Effective Synthesis for Linear Array Thinning," Radar 2021.',
    '"Sparse Array Optimization with Low-Bit Phase Quantization," 2025.',
    '"Pattern Synthesis Using Mayfly Algorithm," IEEE Access, 2021.',
    '"Thinned Planar Arrays Based on Precoded Subarray Structures," 2022.',
    '"High Gain Optimization for Sparse Planar Array," 2022.',
    '"Hybrid Unconstrained-Heuristic for Sparse Planar Arrays," 2022.',
    'Y. Buchris et al., "Frequency-Invariant Beamformers with Sparse Concentric Circular Arrays," WASPAA 2023.',
    '"Window Beamformer for Sparse Concentric Circular Array," 2021.',
    'Y. Wang et al., "Sparse Circular Arrays with Low Sidelobe Levels," CTISC 2022.',
    '"Greedy Design of Circular Concentric Arrays for Broadband MVDR," 2024.',
    'G. Yang et al., "Optimized Design of Sparse Spherical Conformal Array," APCAP 2024.',
    '"Sparse Design of Navigation Conformal Array," 2025.',
    '"Low Sidelobe Optimization of Conformal Sparse Array using Active Pattern," 2024.',
    'G. Liu et al., "Sparse Conformal Array Synthesis Based on Multiagent GA," APCAP 2022.',
    '"Sparse Conformal Multi-Faceted Array," 2023.',
    '"Sparse Cylindrical Arrays Based on Sobol Sequence Sampling," 2022.',
    '"Sparse Cylindrical Arrays with Van der Corput Sequence," 2021.',
    'C. Liu et al., "Multi-Objective Hybrid for Sparse Conformal Array," ICCCA 2025.',
    # Wanfang refs
    '林晗等. 基于改进型灰狼优化算法的直线阵列稀布. 安徽理工大学学报, 2026.',
    '文茜等. 一种基于NSGA-II的稀布平面阵列方向图优化算法. 陕西理工大学学报, 2025.',
    '杜林倩等. 共形天线阵列的稀疏优化. 通信技术, 2021.',
    '国强等. 改进算术优化算法用于稀布平面阵列综合. 西安电子科技大学学报, 2023.',
    '张骆怡等. 大规模平面阵列稀疏优化技术综述. 空间电子技术, 2022.',
    '基于多目标粒子群算法的稀布直线阵优化. 万方.',
    '平面稀布天线阵列的优化算法. 万方.',
    '稀布同心圆环阵列的半径优化方法. 万方.',
    '稀布同心圆环阵列的优化设计. 万方.',
    '基于改进型灰狼优化算法和窗函数加权的稀布矩形平面阵列天线综合. 万方.',
]
for i, ref in enumerate(refs, 1):
    add_reference(doc, i, ref)

add_page_numbers(doc)
doc.save(OUT)
print(f'Done: {OUT} ({os.path.getsize(OUT)} bytes)')
