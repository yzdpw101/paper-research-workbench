# -*- coding: utf-8 -*-
"""
生成《阵列天线稀布优化技术综述》DOCX — 增强版 v2
含：Word数学公式、交叉引用、嵌入的IEEE图片、文献书签
"""
import sys, os
sys.path.insert(0, r'E:/Desktop/稀布综述')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx_formatter import (
    setup_page, add_title, add_abstract_heading, add_abstract_body, add_keywords,
    add_heading1, add_heading2, add_heading3, add_body, add_equation,
    add_figure_with_caption, add_figure_placeholder, add_table_caption, add_table,
    add_reference, add_citation_marker, add_page_numbers, add_seq_field, add_ref_field,
    add_bookmark_around_run, add_bookmark_paragraph, set_run_font
)

TARGET_DIR = r'E:/Desktop/稀布综述'
FIG_DIR = r'E:/Desktop/稀布综述/figures'

def main():
    doc = Document()
    setup_page(doc)

    add_title(doc, "阵列天线稀布优化技术综述")
    
    add_abstract_heading(doc)
    add_abstract_body(doc,
        "阵列天线的稀布（稀疏）优化技术是降低天线系统成本、减轻重量、简化馈电网络的重要手段，"
        "同时可在不显著降低方向图性能的前提下有效抑制旁瓣电平。本文系统综述了2020-2026年间阵列天线稀布优化领域的研究进展，"
        "覆盖线阵、平面阵、圆环阵、多圆环同心环阵和共形阵五种典型阵型。"
        "从优化方法角度，将现有研究分为群体智能算法、凸优化与压缩感知、机器学习方法、"
        "以及解析与混合方法四大类，详细分析各类方法的核心原理、适用场景和典型性能。"
        "最后讨论了工程实现中的互耦效应、幅相误差量化、宽带特性等实际问题，并展望了未来发展方向。")
    add_keywords(doc, "阵列天线；稀布优化；稀疏阵列；旁瓣电平抑制；群体智能算法；凸优化；机器学习；共形阵")

    # ======================== 1 引言 ========================
    add_heading1(doc, "1  引言")
    add_body(doc,
        "阵列天线广泛应用于雷达、通信、电子侦察和射电天文等领域，其性能与阵元数量、排列方式和激励分布密切相关。"
        "传统均匀间隔阵列虽然结构简单、分析方便，但在给定孔径下需要大量阵元单元，导致系统成本高、重量大、热耗散困难。"
        "稀布（稀疏）阵列天线技术通过优化阵元的位置和/或激励幅度，在保持阵列孔径不变的前提下减少有源阵元数量，"
        "从而在不显著恶化方向图性能（特别是峰值旁瓣电平PSLL）的条件下实现系统轻量化和低成本化。")
    add_body(doc,
        "近年来，随着第五代/第六代移动通信、星载相控阵、毫米波雷达和卫星通信等应用的快速发展，"
        "阵列天线的稀布优化技术受到学术界和工业界的广泛关注。"
        "2020-2026年间，该领域涌现了大量新方法、新思路和新应用，特别是在群体智能算法改进、"
        "凸优化框架与压缩感知理论融合、以及深度学习技术引入等方面取得了显著进展。"
        "本文旨在系统综述该时期阵列天线稀布优化领域的研究进展。"
        "第2节建立稀布优化问题的统一数学框架；第3节从方法论角度分类综述各类优化算法；"
        "第4节按阵型分类讨论各类阵列的稀布优化进展；第5节分析工程约束；第6节展望未来方向。")

    # ======================== 2 问题建模 ========================
    add_heading1(doc, "2  稀布阵列天线问题建模")
    add_heading2(doc, "2.1  阵列方向图函数")
    add_body(doc,
        "考虑一个由N个各向同性阵元组成的任意构型阵列，第n个阵元位于位置向量"
        "r_n = (x_n, y_n, z_n)处。阵列的远场方向图函数可表示为：")
    
    # 公式1：阵列因子通式
    add_equation(doc, 
        r"F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta,\phi) \exp[j(k r_n \cos\psi_n + \phi_n)]")
    add_body(doc,
        "其中A_n和φ_n分别为第n个阵元的激励幅度和相位，f_n(θ,φ)为单元方向图，k=2π/λ为波数。"
        "对于等幅同相激励的简化情形，阵列因子简化为：")
    
    # 公式2：简化阵列因子
    add_equation(doc,
        r"\mathrm{AF}(\theta,\phi)=\sum_{n=1}^{N} \exp[jk(x_n\sin\theta\cos\phi + y_n\sin\theta\sin\phi + z_n\cos\theta)]")

    add_heading2(doc, "2.2  各类阵型的阵列因子")
    add_heading3(doc, "2.2.1  线阵")
    add_body(doc, "设N个阵元沿x轴分布在[-L/2, L/2]区间内，阵列因子为：")
    add_equation(doc, r"\mathrm{AF}(\theta)=\sum_{n=1}^{N} \exp(jk x_n \sin\theta)")

    add_heading3(doc, "2.2.2  平面阵")
    add_body(doc,
        "平面阵是最具工程实用价值的构型。设M×N个候选位置按矩形栅格排列，"
        "稀布优化在候选位置中选择K个有源阵元，阵列因子为：")
    add_equation(doc,
        r"\mathrm{AF}(\theta,\phi)=\sum_{m=1}^{M}\sum_{n=1}^{N} I_{mn} \exp[jk(x_m\sin\theta\cos\phi + y_n\sin\theta\sin\phi)]")
    add_body(doc, "其中I_{mn}∈{0,1}为开关变量，表示对应位置阵元的存在与否。")

    add_heading3(doc, "2.2.3  圆环阵与多圆环同心环阵")
    add_body(doc,
        "圆环阵将阵元分布在半径为R的圆周上。多圆环同心环阵由Q个不同半径的圆环组成，阵列因子为：")
    add_equation(doc,
        r"\mathrm{AF}(\theta,\phi)=\sum_{q=1}^{Q}\sum_{n=1}^{N_q} \exp[jk R_q (\sin\theta\cos(\phi-\phi_{qn}))]")

    add_heading3(doc, "2.2.4  共形阵")
    add_body(doc,
        "共形阵将阵元安装在非平面载体表面。由于各阵元指向方向不同，阵列方向图计算时必须计入单元方向图的指向依赖性：")
    add_equation(doc,
        r"F(\theta,\phi)=\sum_{n=1}^{N} A_n f_n(\theta-\theta_n,\phi-\phi_n) \exp[jk\mathbf{r}_n\cdot\hat{\mathbf{r}}]")

    add_heading2(doc, "2.3  优化问题形式化")
    add_body(doc,
        "阵列稀布优化的标准形式为约束非线性优化问题：最小化峰值旁瓣电平PSLL，"
        "约束条件包括阵元数上限、孔径尺寸、最小间距、激励动态范围等。"
        "目标函数可表示为：")
    add_equation(doc,
        r"\min_{\{\mathbf{r}_n,A_n\}} \max_{(\theta,\phi)\in\Theta_{\mathrm{SL}}} |\mathrm{AF}(\theta,\phi)| \quad \text{s.t.} \quad N\leq N_{\max},\; d_{\min}\leq\|\mathbf{r}_i-\mathbf{r}_j\|")

    # ======================== 3 优化方法 ========================
    add_heading1(doc, "3  稀布阵列优化方法综述")
    add_body(doc,
        "现有的稀布阵列优化方法可分为四类：群体智能算法、凸优化与压缩感知、机器学习方法、以及解析与混合方法。")

    add_heading2(doc, "3.1  群体智能算法")
    add_heading3(doc, "3.1.1  遗传算法")
    add_body(doc,
        "遗传算法（GA）是稀布优化中最经典的方法。文献[1]使用二进制遗传算法（BGA）对线阵进行稀疏优化。"
        "文献[2]提出基于量子选择机制的GA用于平面阵稀疏。文献[3]采用自适应多点变异GA优化大规模稀疏阵列。"
        "文献[4]研究了线阵稀疏对方向图参数的影响规律。")

    add_heading3(doc, "3.1.2  粒子群优化")
    add_body(doc,
        "粒子群优化（PSO）及其二进制变体（BPSO）广泛应用于阵列稀疏。"
        "文献[5]采用BPSO实现了平面阵的稀疏综合。文献[6]提出了多策略非线性多目标PSO用于稀疏平面阵综合。"
        "文献[7]将模拟退火融入PSO设计了最优MIMO稀疏阵列。")

    add_heading3(doc, "3.1.3  其他新兴算法")
    add_body(doc,
        "近年来各种新兴群智能算法被引入阵列稀布优化。文献[8]提出改进二进制灰狼优化（GWO）算法。"
        "文献[9]使用二进制蝙蝠群算法（BBA）进行线阵稀疏。文献[10]将布谷鸟搜索（CS）与凸优化相结合。"
        "文献[11]提出改进蜣螂优化算法用于平面阵稀布。文献[12]采用改进蛇优化算法实现共形阵稀疏波束形成。")

    # ===== 图1：群体智能算法流程 =====
    fig1_path = os.path.join(FIG_DIR, "算法/9753378-fig-3-source.gif")
    if os.path.isfile(fig1_path):
        add_figure_with_caption(doc, fig1_path,
            "基于CNN的阵列天线稀疏综合流程（文献[13]）")
    else:
        add_figure_placeholder(doc, "群体智能算法在阵列稀布优化中的应用流程示意图",
                              "综合文献[1][5][8]方法流程")

    add_heading2(doc, "3.2  凸优化与压缩感知方法")
    add_body(doc,
        "凸优化方法具有严格的数学收敛保证。文献[14]通过迭代凸优化对稀疏阵列进行方向图综合。"
        "文献[15]将交替方向乘子法（ADMM）应用于稀疏阵列方向图综合。"
        "文献[16]采用共识ADMM框架实现可重构稀疏阵列的多波束综合。"
        "文献[17]提出重加权l1范数最小化算法。")
    add_body(doc,
        "压缩感知为阵列稀疏优化提供了新的途径。文献[18]基于离网压缩感知实现了低复杂度稀疏阵列综合。"
        "文献[19]将贝叶斯压缩感知（BCS）与阵列扩张技术结合。"
        "文献[20]将分支定界法与凸优化结合用于大规模稀疏阵列卫星通信。")

    add_heading2(doc, "3.3  机器学习方法")
    add_body(doc,
        "文献[13]使用卷积神经网络（CNN）实现了阵列天线的稀疏综合。文献[21]提出了基于深度神经网络（DNN）的线阵稀疏综合方法。"
        "文献[22]提出了自适应概率密度锥削学习方法，通过强化学习框架优化稀疏概率分布。")

    add_heading2(doc, "3.4  解析与混合方法")
    add_body(doc,
        "迭代傅里叶变换（IFT）是最具代表性的解析方法。文献[23]使用两级IFT实现稀疏平面阵综合。"
        "文献[24]将JAYA算法与IFT混合用于线阵稀疏。文献[25]提出了OMP-Broyden混合算法。"
        "文献[26]采用量子-经典混合进化优化实现阵列稀疏。")

    # ===== 表1：方法对比 =====
    add_table_caption(doc, "各类稀布优化方法对比")
    add_table(doc,
        ["方法类别", "代表算法", "优势", "局限性", "适用阵型"],
        [
            ["群体智能", "GA/BPSO/GWO/CS", "无梯度、全局搜索", "计算量大", "全部阵型"],
            ["凸优化/CS", "ADMM/SOCP", "收敛保证", "约束敏感", "线阵/平面阵"],
            ["机器学习", "CNN/DNN/RL", "推理快", "泛化待验证", "线阵/平面阵"],
            ["解析方法", "IFT/矩阵束", "极快", "灵活性差", "线阵/平面阵"],
            ["混合方法", "JAYA-IFT/Broyden", "兼顾效率与搜索", "设计复杂", "线阵/平面阵"],
        ])

    # ======================== 4 各阵型 ========================
    add_heading1(doc, "4  各阵型稀布优化研究进展")
    add_heading2(doc, "4.1  稀布线阵")
    add_body(doc,
        "线阵的稀布优化研究最为成熟。文献[1-4][8-10]等从不同算法角度降低线阵PSLL。"
        "文献[27]提出了基于点扩散函数（PSF）的新适应度函数。文献[28]针对太赫兹频段的大规模线阵进行GA稀疏设计。"
        "文献[29]提出了一种有效的线阵稀疏综合技术。文献[30]将低比特相位量化约束引入线阵稀疏优化。"
        "文献[31]综合比较了多种群智能算法在线阵稀疏中的性能。")

    # ===== 图2：线阵稀布优化结果 =====
    fig2_path = os.path.join(FIG_DIR, "线阵/9964507-fig-4-source.gif")
    if os.path.isfile(fig2_path):
        add_figure_with_caption(doc, fig2_path,
            "BGA优化后稀布线阵的方向图（文献[1]）")

    add_heading2(doc, "4.2  稀布平面阵")
    add_body(doc,
        "平面阵的稀布优化因二维自由度更高而更具挑战性。文献[5][14][23]分别从智能算法、压缩感知和IFT角度解决平面阵稀疏问题。"
        "文献[32]提出了基于预编码子阵结构的平面阵稀疏方法。文献[33]在等幅激励条件下优化稀疏平面阵的增益。"
        "文献[34]采用混合无约束-启发式方法实现多种稀疏平面阵的综合。")

    # ===== 图3：平面阵方向图 =====
    fig3_path = os.path.join(FIG_DIR, "平面阵/yang4-3088492.gif")
    if os.path.isfile(fig3_path):
        add_figure_with_caption(doc, fig3_path,
            "迭代凸优化稀疏平面阵的3D方向图（文献[14]）")

    add_heading2(doc, "4.3  圆环阵与多圆环同心环阵")
    add_body(doc,
        "圆环阵的稀布优化近年受到越来越多的关注。文献[35]设计了频率不变波束形成器的稀疏同心圆环阵。"
        "文献[36]提出了加窗波束形成器用于稀疏同心圆环阵。文献[37]直接研究低旁瓣稀疏圆环阵的设计方法。"
        "文献[38]利用贪婪算法实现了宽带MVDR波束形成的同心圆环阵稀疏设计。")

    # ===== 图4：圆环阵 =====
    fig4_path = os.path.join(FIG_DIR, "圆环阵/wang1-p5-wang.gif")
    if os.path.isfile(fig4_path):
        add_figure_with_caption(doc, fig4_path,
            "稀疏圆环阵阵元分布与方向图（文献[37]）")

    add_heading2(doc, "4.4  共形阵")
    add_body(doc,
        "共形阵的稀布优化是领域前沿。文献[39]基于智能算法实现稀疏球面共形阵优化。文献[40]针对导航共形阵提出稀疏设计。"
        "文献[41]利用有源单元方向图（AEP）方法实现低旁瓣共形稀疏阵。文献[42]基于多智能体GA实现稀疏共形阵综合。"
        "文献[43]设计了基于金字塔/梯形子阵的稀疏共形多面阵。")
    add_body(doc,
        "在稀疏柱面阵方面，文献[44][45]分别基于Sobol序列和Van der Corput序列实现低偏差采样。"
        "文献[46]通过多目标混合方法优化稀疏共形阵的方向图。")

    # ===== 图5：共形阵 =====
    fig5_path = os.path.join(FIG_DIR, "共形阵/10069223-fig-2-source.gif")
    if os.path.isfile(fig5_path):
        add_figure_with_caption(doc, fig5_path,
            "多智能体GA优化的稀疏共形阵阵元分布（文献[42]）")
    else:
        add_figure_placeholder(doc, "典型共形阵构型：柱面阵/球面阵/锥台阵",
                              "综合文献[39][42][44]")

    # ======================== 5 工程约束 ========================
    add_heading1(doc, "5  工程约束与讨论")
    add_heading2(doc, "5.1  互耦效应")
    add_body(doc,
        "稀布阵列的阵元间距非均匀性导致互耦效应复杂化。文献[41]通过引入AEP来考虑互耦影响。"
        "将互耦效应纳入稀布优化模型是重要的开放问题。")

    add_heading2(doc, "5.2  幅相误差与量化约束")
    add_body(doc,
        "文献[30]将低比特相位量化约束嵌入优化模型。文献[17]采用混合整数规划处理离散相位约束。")

    add_heading2(doc, "5.3  宽带特性")
    add_body(doc,
        "文献[35]针对宽带波束形成设计频率不变稀疏圆环阵。文献[38]设计宽带MVDR波束形成的同心圆环阵。")

    # ======================== 6 展望 ========================
    add_heading1(doc, "6  未来展望")
    add_body(doc, "阵列天线稀布优化技术在未来将呈现以下发展趋势：")
    add_body(doc,
        "（1）AI驱动的端到端设计：将深度学习与电磁仿真深度耦合，实现从构型设计到性能验证的全流程自动化。"
        "（2）亚波长与可重构稀布阵列：结合超表面、MEMS开关等可调元件实现波束可重构。"
        "（3）MIMO雷达与通信一体化中的稀布设计：大规模MIMO阵列需同时考虑通信和感知的联合性能。"
        "（4）多物理场协同优化：在天线结构、热管理和电磁性能之间建立联合优化框架。"
        "（5）量子计算在阵列优化中的应用：量子-经典混合算法有望为大规模阵列优化提供新路径。")

    # ======================== 参考文献 ========================
    add_heading1(doc, "参考文献")
    refs = [
        "M. M. Kamal et al., \"Optimization of Linear Antenna Array Thinning using Binary Genetic Algorithm (BGA),\" ICCCI 2022.",
        "\"Planar Array Thinning by Genetic Algorithm with Quantum Selection,\" 2025.",
        "\"Sparse Antenna Array Synthesis via Adaptive Multipoint Mutation Genetic Algorithm,\" 2025.",
        "A. S. Karasev et al., \"Effect of Linear Antenna Array Thinning on Its Directional Pattern Parameters,\" EDM 2022.",
        "S. Vankayalapati et al., \"Application of Binary Particle Swarm Optimization for Thinned Planar Array Synthesis,\" ICACCS 2022.",
        "\"Multi-Strategy Nonlinear Multi-Objective PSO for Sparse Planar Array Synthesis,\" 2025.",
        "\"Optimal MIMO Sparse Array Design Based on Simulated Annealing PSO,\" 2022.",
        "L. Su et al., \"Linear Antenna Array Thinning Based on Improved Binary GWO,\" ICMMT 2025.",
        "S. Vankayalapati et al., \"Binary bat swarm algorithm for thinning antenna elements in linear array,\" ICCCI 2023.",
        "R.-Q. Wang et al., \"Synthesis of Sparse Linear Arrays Using Hybrid Cuckoo Search with Convex Programming,\" IEEE AWPL, 2020.",
        "\"Synthesis of Planar Sparse Arrays Based on Improved Dung Beetle Optimization Algorithm,\" 2025.",
        "\"Improved Snake Optimization Algorithm for Sparse Conformal Array Beamforming,\" 2024.",
        "S. Shen et al., \"Synthesis of Thinned Array Antenna with Convolutional Neural Network,\" ISAPE 2021.",
        "\"Synthesis of Sparse Antenna Arrays Subject to Constraint on Directivity via Iterative Convex Optimization,\" 2021.",
        "\"Sparse Array Beampattern Synthesis via Majorization-Based ADMM,\" 2021.",
        "\"Multibeam Synthesis for Reconfigurable Sparse Array via Consensus PDD Framework,\" 2023.",
        "\"Synthesis of Sparse Arrays With Reweighted l1-norm Minimization,\" 2022.",
        "\"Low-Complexity Sparse Array Synthesis Based on Off-Grid Compressive Sensing,\" 2022.",
        "\"Bayesian Compressive Sensing and Array Dilation for Sparse Linear Arrays,\" 2022.",
        "\"Large-Scale Sparse Array Using Branch and Bound with Convex Optimization,\" 2025.",
        "\"Synthesis of Linear Sparse Array Using DNN-Based Machine-Learning Method,\" 2023.",
        "\"Adaptive Learning of Probability Density Taper for Large Planar Array Thinning,\" 2020.",
        "\"Synthesis of Thinned Planar Array Using Two-Stage Iterative Fourier Transform,\" 2025.",
        "X. Yang et al., \"Linear Array Thinning Using JAYA-Iterative Fourier Technique,\" ACES-China 2023.",
        "\"A Hybrid OMP-Broyden Algorithm for Thinning Antenna Array,\" 2025.",
        "\"Array Thinning by Hybrid Quantum-Classical Evolutionary Optimization,\" 2025.",
        "\"A New Fitness Function for Sparse Linear Array Evaluation Based on PSF,\" 2021.",
        "\"Static Thinning of Large Linear Antenna Arrays Using GA for THz Frequencies,\" 2025.",
        "Z. Zhou et al., \"An Effective and Efficient Synthesis Technique for Linear Array Thinning,\" Radar 2021.",
        "\"Sparse Array Optimization for Beampattern Synthesis with Low-Bit Phase Quantization,\" 2025.",
        "\"Pattern Synthesis of Uniform and Sparse Linear Array Using Mayfly Algorithm,\" IEEE Access, 2021.",
        "\"Synthesis of Thinned Planar Arrays Based on Precoded Subarray Structures,\" 2022.",
        "\"High Gain Optimization with Equal Amplitude Excitation for Sparse Planar Array,\" 2022.",
        "\"Pattern Synthesis for Sparse Planar Arrays by Hybrid Unconstrained-Heuristic Approach,\" 2022.",
        "Y. Buchris et al., \"Frequency-Invariant Beamformers with Sparse Concentric Circular Arrays,\" WASPAA 2023.",
        "\"Window Beamformer for Sparse Concentric Circular Array,\" 2021.",
        "Y. Wang et al., \"Sparse Circular Arrays with Low Sidelobe Levels,\" CTISC 2022.",
        "\"Greedy Design of Circular Concentric Arrays for Broadband MVDR,\" 2024.",
        "G. Yang et al., \"Optimized Design of Sparse Spherical Conformal Array,\" APCAP 2024.",
        "\"Sparse Design of Navigation Conformal Array and Its Beamforming,\" 2025.",
        "\"Low Sidelobe Optimization of Conformal Sparse Array using Active Pattern,\" 2024.",
        "G. Liu et al., \"Sparse Conformal Array Synthesis Based on Multiagent GA,\" APCAP 2022.",
        "\"Sparse Conformal Multi-Faceted Array Based on Pyramid/Trapezoid Subarrays,\" 2023.",
        "\"Sparse Cylindrical Arrays Based on Low-Discrepancy Sobol Sequence Sampling,\" 2022.",
        "\"Sparse Cylindrical Arrays with Van der Corput Sequence Spacing,\" 2021.",
        "C. Liu et al., \"Multi-Objective Hybrid Method for Sparse Conformal Array Directivity,\" ICCCA 2025.",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)

    add_page_numbers(doc)
    output_path = os.path.join(TARGET_DIR, "阵列天线稀布优化技术综述.docx")
    doc.save(output_path)
    print(f"综述已保存至: {output_path}")
    print(f"段落数: {len(doc.paragraphs)} | 参考文献: {len(refs)}")

if __name__ == "__main__":
    main()
